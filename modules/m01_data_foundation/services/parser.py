from __future__ import annotations

import re
import statistics
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedBlock:
    section_path: str
    block_type: str
    block_text: str
    parent_index: int | None = None
    page_no: str | None = None
    start_offset: int | None = None
    end_offset: int | None = None
    metadata_json: dict = field(default_factory=dict)


class DocumentParser:
    """层级文段解析器（README FR-PARSE-002 / FR-PARSE-003）。

    将文档解析为带结构层级的小文段：标题自动推断层级、构建 section_path、
    记录 parent_index（指向所属标题在结果列表中的位置，落库时再解析为 block_id），
    PDF 额外记录 page_no，所有块保留 start/end_offset 与 metadata 便于原文定位。
    """

    # 相邻同类型非标题块合并的字符上限，超过则另起一块，避免整篇合成超大块
    MAX_MERGE_CHARS = 1200

    def parse_document(self, file_path: Path, file_type: str) -> list[ParsedBlock]:
        suffix = file_type.lower() if file_type else file_path.suffix.lower()
        if suffix in {".pdf"}:
            raw = self._read_pdf(file_path)
        elif suffix in {".docx", ".doc"}:
            raw = self._read_docx(file_path)
        else:
            raw = self._read_text(file_path)
        raw = self._merge_consecutive(raw)
        return self._assign_hierarchy(raw)

    # ------------------------------------------------------------------
    # 各格式读取：产出扁平 raw 列表，每项含 text / level / block_type 等
    # ------------------------------------------------------------------
    def _read_text(self, file_path: Path) -> list[dict]:
        text = self._safe_read(file_path)
        return self._md_lines_to_raw(text)

    def _read_pdf(self, file_path: Path) -> list[dict]:
        import fitz

        doc = fitz.open(file_path)
        raw: list[dict] = []
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            page_no = str(page_index + 1)
            blocks = page.get_text("dict").get("blocks", [])
            sizes = [
                span["size"]
                for b in blocks
                if b.get("type") == 0
                for line in b.get("lines", [])
                for span in line.get("spans", [])
            ]
            body_size = statistics.median(sizes) if sizes else 11.0
            for b in blocks:
                if b.get("type") != 0:
                    continue
                spans_texts: list[str] = []
                max_size = 0.0
                bold = False
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        spans_texts.append(span["text"])
                        max_size = max(max_size, span["size"])
                        if span.get("flags", 0) & 16:  # bold
                            bold = True
                block_text = "".join(spans_texts).strip()
                if not block_text:
                    continue
                is_title = (max_size >= body_size * 1.2) or (bold and len(block_text) < 40)
                level = self._pdf_level(max_size, body_size) if is_title else None
                raw.append(
                    {
                        "text": block_text,
                        "level": level,
                        "block_type": "title" if is_title else "paragraph",
                        "page_no": page_no,
                        "metadata": {"font_size": round(max_size, 1)},
                    }
                )
        return raw

    def _read_docx(self, file_path: Path) -> list[dict]:
        import docx

        document = docx.Document(str(file_path))
        raw: list[dict] = []
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            match = re.match(r"^Heading\s+(\d+)", style_name)
            if match:
                raw.append({"text": text, "level": int(match.group(1)), "block_type": "title"})
                continue
            # 中文制度/规程类文档常未套用 Heading 样式，按常见章节/条款模式推断层级
            heading_level = self._docx_heading_level(text)
            if heading_level is not None:
                raw.append({"text": text, "level": heading_level, "block_type": "title"})
            else:
                raw.append({"text": text, "level": None, "block_type": "paragraph"})
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                joined = " | ".join(cells)
                if joined.strip():
                    raw.append({"text": joined, "level": None, "block_type": "table_row"})
        return raw

    @staticmethod
    def _docx_heading_level(text: str) -> int | None:
        """按中文文档常见标题模式推断层级（章=1 / 节·部分=2 / 条=3 / 附件=2）。"""
        if re.match(r"^第[一二三四五六七八九十百千零两]+\s*章", text):
            return 1
        if re.match(r"^第[一二三四五六七八九十百千零两]+\s*[节部分]", text):
            return 2
        if re.match(r"^第[一二三四五六七八九十百千零两]+\s*条", text):
            return 3
        if re.match(r"^(附件|附录)\s*[0-9一二三四五六七八九十]*", text):
            return 2
        return None

    # ------------------------------------------------------------------
    # 层级构建
    # ------------------------------------------------------------------
    def _assign_hierarchy(self, raw: list[dict]) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        # 栈保存 (在 blocks 中的位置, level)，栈顶为当前最近的标题
        stack: list[tuple[int, int]] = []
        for item in raw:
            level = item.get("level")
            # 标题先出栈同级/更高级祖先，再计算其所属层级
            if level is not None:
                while stack and stack[-1][1] >= level:
                    stack.pop()
            ancestor_texts = [blocks[i].block_text for i, _ in stack]
            section_path = " / ".join(ancestor_texts) if ancestor_texts else "未分类"
            parent_index = stack[-1][0] if stack else None
            block = ParsedBlock(
                section_path=section_path,
                block_type=item.get("block_type", "paragraph"),
                block_text=item["text"],
                parent_index=parent_index,
                page_no=item.get("page_no"),
                start_offset=item.get("start"),
                end_offset=item.get("end"),
                metadata_json=item.get("metadata", {}),
            )
            blocks.append(block)
            if level is not None:
                stack.append((len(blocks) - 1, level))
        return blocks

    # ------------------------------------------------------------------
    # 工具
    # ------------------------------------------------------------------
    @staticmethod
    def _md_lines_to_raw(text: str) -> list[dict]:
        raw: list[dict] = []
        offset = 0
        for line in text.split("\n"):
            line_start = offset
            offset += len(line) + 1
            stripped = line.strip()
            if not stripped:
                continue
            match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if match:
                raw.append(
                    {
                        "text": match.group(2).strip(),
                        "level": len(match.group(1)),
                        "block_type": "title",
                        "start": line_start,
                        "end": line_start + len(line),
                    }
                )
            else:
                list_match = re.match(r"^([-*+]|\d+[.)])\s+", stripped)
                raw.append(
                    {
                        "text": stripped,
                        "level": None,
                        "block_type": "list_item" if list_match else "paragraph",
                        "start": line_start,
                        "end": line_start + len(line),
                    }
                )
        return raw

    @staticmethod
    def _merge_consecutive(raw: list[dict]) -> list[dict]:
        """合并相邻的非标题块（同类型），避免一行一段造成过碎；但受 MAX_MERGE_CHARS 限制。"""
        merged: list[dict] = []
        for item in raw:
            if item.get("level") is not None:
                merged.append(item)
                continue
            last = merged[-1] if merged else None
            if (
                last is not None
                and last.get("level") is None
                and last.get("block_type") == item.get("block_type")
                and len(last["text"]) + len(item["text"]) <= DocumentParser.MAX_MERGE_CHARS
            ):
                last["text"] = f"{last['text']}\n{item['text']}"
                last["end"] = item.get("end")
            else:
                merged.append(item)
        return merged

    @staticmethod
    def _pdf_level(max_size: float, body_size: float) -> int:
        ratio = max_size / body_size if body_size else 1.0
        if ratio >= 1.6:
            return 1
        if ratio >= 1.3:
            return 2
        return 3

    @staticmethod
    def _safe_read(file_path: Path) -> str:
        for enc in ("utf-8", "utf-8-sig", "gb18030", "latin-1"):
            try:
                return Path(file_path).read_text(encoding=enc)
            except (UnicodeDecodeError, UnicodeError):
                continue
        return Path(file_path).read_text(encoding="utf-8", errors="replace")
